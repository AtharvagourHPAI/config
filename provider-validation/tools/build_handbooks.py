"""Generate the two project handbooks as Word (.docx) documents.

Outputs (written to ``docs/`` at the repo root):
  1. Local_Setup_Guide.docx  - step-by-step guide to run the app locally.
  2. API_Handbook.docx        - reference for the FastAPI REST layer.

Run with:  .venv\\Scripts\\python.exe tools/build_handbooks.py
"""

from __future__ import annotations

import pathlib

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = pathlib.Path(__file__).resolve().parent.parent
DOCS = ROOT / "docs"

BRAND = RGBColor(0x1F, 0x4E, 0x79)      # deep blue
ACCENT = RGBColor(0x2E, 0x75, 0xB6)     # lighter blue
CODE_BG = "F2F2F2"
HEADER_BG = "1F4E79"
ZEBRA_BG = "EAF1F8"
MONO = "Consolas"
BODY = "Calibri"


# --------------------------------------------------------------------------- #
# Low-level helpers
# --------------------------------------------------------------------------- #
def _shade(cell, fill: str) -> None:
    """Apply a solid background fill to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.makeelement(qn("w:shd"), {})
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_text(cell, text, *, bold=False, color=None, mono=False, size=10):
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = MONO if mono else BODY
    if color is not None:
        run.font.color.rgb = color


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BRAND if level <= 1 else ACCENT
    return h


def add_body(doc, text, *, size=11, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = BODY
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(item, tuple):
            lead, rest = item
            r = p.add_run(lead)
            r.bold = True
            r.font.name = MONO
            p.add_run(rest)
        else:
            p.add_run(item)
        for run in p.runs:
            run.font.size = Pt(11)


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text).font.size = Pt(11)
    return p


def add_code(doc, code: str):
    """Render a monospaced, shaded code block."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _shade(cell, CODE_BG)
    cell.text = ""
    for i, line in enumerate(code.strip("\n").split("\n")):
        para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        run = para.add_run(line if line else " ")
        run.font.name = MONO
        run.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, head in enumerate(headers):
        _shade(hdr[i], HEADER_BG)
        _set_cell_text(hdr[i], head, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=10)
    for r, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            mono = i == 0 and len(headers) > 1
            _set_cell_text(cells[i], val, mono=mono, size=10)
            if r % 2 == 1:
                _shade(cells[i], ZEBRA_BG)
    return table


def add_cover(doc, title, subtitle, meta_lines):
    for _ in range(3):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(title)
    run.bold = True
    run.font.size = Pt(30)
    run.font.color.rgb = BRAND
    run.font.name = BODY

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = s.add_run(subtitle)
    sr.font.size = Pt(15)
    sr.font.color.rgb = ACCENT

    for _ in range(2):
        doc.add_paragraph()
    for line in meta_lines:
        m = doc.add_paragraph()
        m.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mr = m.add_run(line)
        mr.font.size = Pt(11)
        mr.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    doc.add_page_break()


def base_document() -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = BODY
    style.font.size = Pt(11)
    return doc


# --------------------------------------------------------------------------- #
# Document 1: Local Setup Guide
# --------------------------------------------------------------------------- #
def build_local_setup():
    doc = base_document()
    add_cover(
        doc,
        "Step-by-Step Local Setup Guide",
        "Provider Contract Change Validation Engine (CMS-855I)",
        [
            "Run the REST API, the Streamlit reviewer UI, and the test suite locally",
            "Audience: developers and reviewers setting up the project for the first time",
            "Version 1.0",
        ],
    )

    add_heading(doc, "1. Overview", 1)
    add_body(
        doc,
        "The Provider Contract Change Validation Engine is a deterministic, "
        "explainable rule engine for Medicare CMS-855I provider-enrollment changes. "
        "It ships with three runnable surfaces, all backed by the same decision core "
        "in engine/:",
    )
    add_bullets(
        doc,
        [
            ("Test suite \u2014 ", "pytest runs the loaders, rules, decision, regression, and API tests."),
            ("REST API \u2014 ", "a FastAPI service exposing the engine over HTTP (Swagger docs at /docs)."),
            ("Streamlit UI \u2014 ", "a reviewer dashboard to upload change requests and inspect decision traces."),
        ],
    )
    add_body(
        doc,
        "This guide takes you from a fresh checkout to all three running on your "
        "machine. Commands are shown for Windows PowerShell (the project's primary "
        "environment); macOS/Linux equivalents are noted where they differ.",
    )

    add_heading(doc, "2. Prerequisites", 1)
    add_table(
        doc,
        ["Requirement", "Version / Notes"],
        [
            ["Python", "3.11 or newer (developed and verified on 3.14)"],
            ["pip", "Bundled with Python; used to install dependencies"],
            ["OS shell", "Windows PowerShell, or bash/zsh on macOS/Linux"],
            ["Disk", "A few hundred MB for the virtual environment"],
            ["Network", "Needed only once, to download Python packages"],
        ],
    )
    add_body(doc, "Verify Python is available before continuing:")
    add_code(doc, "python --version")
    add_body(
        doc,
        "If python is not recognized on Windows, disable the Microsoft Store "
        "\u201cApp execution alias\u201d for Python or use the full path to your "
        "Python install (for example py -3.11).",
    )

    add_heading(doc, "3. Get the project", 1)
    add_body(
        doc,
        "Open a terminal in the project root. All commands below assume you are "
        "inside the provider-validation folder:",
    )
    add_code(
        doc,
        'cd "C:\\Users\\<you>\\...\\Config AI\\provider-validation"',
    )

    add_heading(doc, "4. Create and activate a virtual environment", 1)
    add_body(
        doc,
        "A virtual environment keeps the project's dependencies isolated from your "
        "system Python.",
    )
    add_heading(doc, "Windows (PowerShell)", 2)
    add_code(
        doc,
        "python -m venv .venv\n"
        ".\\.venv\\Scripts\\Activate.ps1",
    )
    add_body(
        doc,
        "If activation is blocked by the execution policy, run once: "
        "Set-ExecutionPolicy -Scope Process -ExecutionPolicy Bypass.",
    )
    add_heading(doc, "macOS / Linux", 2)
    add_code(
        doc,
        "python3 -m venv .venv\n"
        "source .venv/bin/activate",
    )
    add_body(
        doc,
        "Once active, your prompt is prefixed with (.venv). Use this same "
        "environment for every command that follows.",
    )

    add_heading(doc, "5. Install dependencies", 1)
    add_body(
        doc,
        "Install everything the project needs from requirements.txt:",
    )
    add_code(doc, "pip install -r requirements.txt")
    add_body(doc, "This installs, among others:")
    add_table(
        doc,
        ["Package", "Used for"],
        [
            ["pandas, openpyxl", "Reading the Excel reference / change-request workbooks"],
            ["pydantic", "Typed data models and validation"],
            ["pyyaml", "Loading the editable rulebook (config/rules.yaml)"],
            ["fastapi, uvicorn", "The REST API and its dev server"],
            ["streamlit, altair", "The reviewer UI and its charts"],
            ["pytest, httpx", "The test suite and API test client"],
        ],
    )

    add_heading(doc, "6. Verify the install with the test suite", 1)
    add_body(
        doc,
        "Before running anything, confirm the engine is healthy. From the project "
        "root, run:",
    )
    add_code(doc, "pytest")
    add_body(
        doc,
        "You should see all tests pass (46 tests, including the full regression "
        "that reproduces every labeled outcome and the synthetic REJECT fixtures). "
        "A green run confirms data loading, the rulebook, decision precedence, and "
        "the API all work end to end.",
    )

    add_heading(doc, "7. Run the REST API", 1)
    add_body(doc, "Start the FastAPI service with uvicorn:")
    add_code(doc, "uvicorn api.main:app --reload")
    add_body(doc, "The server starts on http://127.0.0.1:8000. Once it is up:")
    add_table(
        doc,
        ["URL", "What you get"],
        [
            ["http://127.0.0.1:8000/docs", "Interactive Swagger UI \u2014 try every endpoint"],
            ["http://127.0.0.1:8000/redoc", "ReDoc reference documentation"],
            ["http://127.0.0.1:8000/health", "Health check + loaded contract/change/rule counts"],
            ["http://127.0.0.1:8000/openapi.json", "Machine-readable OpenAPI schema"],
        ],
    )
    add_body(
        doc,
        "The --reload flag restarts the server automatically when you edit a "
        "source file \u2014 ideal for development. Press CTRL+C to stop it.",
    )
    add_body(
        doc,
        "Quick smoke test (in a second terminal) confirming the API is live:",
    )
    add_code(doc, "curl http://127.0.0.1:8000/health")
    add_body(
        doc,
        "See the companion API Handbook for the full endpoint reference and example "
        "request/response payloads.",
    )

    add_heading(doc, "8. Run the Streamlit reviewer UI", 1)
    add_body(
        doc,
        "In a separate terminal (with the virtual environment activated), launch "
        "the UI:",
    )
    add_code(doc, "streamlit run app/main.py")
    add_body(
        doc,
        "Streamlit opens a browser tab (typically http://localhost:8501). The UI "
        "has three tabs:",
    )
    add_bullets(
        doc,
        [
            ("Input \u2014 ", "loads the bundled reference data and lets you upload a change-request workbook (or use the sample), then click Process to score it."),
            ("Single decision \u2014 ", "pick one change request and inspect its full trace: outcome, contract action, fired tags, and every evaluated rule."),
            ("Batch report \u2014 ", "overall accuracy, the expected-vs-actual outcome distribution, a confusion matrix, and per-category accuracy."),
        ],
    )

    add_heading(doc, "9. Optional configuration", 1)
    add_body(doc, "Two environment variables tune the API at startup:")
    add_table(
        doc,
        ["Variable", "Effect"],
        [
            ["DATA_PATH", "Directory holding reference_data.xlsx and contract_change_requests.xlsx. Unset = bundled data/."],
            ["ALLOWED_ORIGINS", "Comma-separated CORS origins for the API. Unset = none (no cross-origin browser calls)."],
        ],
    )
    add_body(doc, "Example (PowerShell), pointing the API at a custom data folder:")
    add_code(
        doc,
        '$env:DATA_PATH = "C:\\data\\custom"\n'
        "uvicorn api.main:app --reload",
    )
    add_body(
        doc,
        "The editable rulebook lives in config/rules.yaml and the outcome\u2192action "
        "mapping in config/outcomes.yaml. You can retune rule scope, failure "
        "outcomes, and vocabularies there without touching code.",
    )

    add_heading(doc, "10. Troubleshooting", 1)
    add_table(
        doc,
        ["Symptom", "Fix"],
        [
            ["'python' is not recognized", "Disable the Windows Store Python alias, or use py -3.11 / the full path."],
            ["Activate.ps1 is blocked", "Run Set-ExecutionPolicy -Scope Process -ExecutionPolicy Bypass, then re-activate."],
            ["ModuleNotFoundError: No module named 'engine'", "Run commands from the provider-validation root, not a subfolder."],
            ["Port 8000 already in use", "Stop the other process, or run uvicorn api.main:app --port 8001."],
            ["Port 8501 already in use", "Run streamlit run app/main.py --server.port 8502."],
            ["Tests fail after editing rules.yaml", "Revert the change or update the expected outcomes; the regression asserts exact label parity."],
        ],
    )

    add_heading(doc, "11. Quick reference", 1)
    add_body(doc, "The full local workflow, start to finish:")
    add_code(
        doc,
        "# 1. create + activate venv\n"
        "python -m venv .venv\n"
        ".\\.venv\\Scripts\\Activate.ps1\n\n"
        "# 2. install deps\n"
        "pip install -r requirements.txt\n\n"
        "# 3. verify\n"
        "pytest\n\n"
        "# 4. run the API (terminal A)\n"
        "uvicorn api.main:app --reload\n\n"
        "# 5. run the UI (terminal B)\n"
        "streamlit run app/main.py",
    )

    out = DOCS / "Local_Setup_Guide.docx"
    doc.save(out)
    return out


# --------------------------------------------------------------------------- #
# Document 2: API Handbook
# --------------------------------------------------------------------------- #
def build_api_handbook():
    doc = base_document()
    add_cover(
        doc,
        "API Handbook",
        "Provider Contract Change Validation API (CMS-855I)",
        [
            "REST reference for the FastAPI layer over the deterministic decision engine",
            "Base URL: http://127.0.0.1:8000   \u2022   API version 1.0.0",
            "Audience: integrators calling the validation service over HTTP",
        ],
    )

    add_heading(doc, "1. Introduction", 1)
    add_body(
        doc,
        "The API is a thin HTTP adapter over the deterministic CMS-855I decision "
        "engine. It contains no decision, tag, or rule logic of its own \u2014 every "
        "endpoint imports engine.decision_engine.decide() and serializes the result. "
        "The reference workbook (contracts + rulebook) is loaded exactly once at "
        "startup and shared across requests.",
    )
    add_body(
        doc,
        "Given a change request and the provider's current contract, the engine "
        "returns exactly one of five outcomes plus a downstream contract action and "
        "a full, explainable decision trace.",
    )

    add_heading(doc, "2. Key concepts", 1)
    add_heading(doc, "2.1 Outcomes", 2)
    add_table(
        doc,
        ["Outcome", "Meaning"],
        [
            ["APPROVE", "Change passes all rules; apply it to the contract."],
            ["DEVELOP", "Additional documentation/verification needed before deciding."],
            ["DENY", "Hard block (e.g. exclusion/debarment, incompatible specialty)."],
            ["REJECT", "Procedural/transport defect (wrong CMS form, missing signature)."],
            ["INITIAL_ENROLLMENT_REQUIRED", "Out-of-state practice-location add needs a new enrollment."],
        ],
    )
    add_heading(doc, "2.2 Status-code philosophy", 2)
    add_body(
        doc,
        "Business outcomes are all successful evaluations and return HTTP 200 \u2014 "
        "even a DENY or REJECT. The 4xx range is reserved strictly for transport "
        "problems.",
    )
    add_table(
        doc,
        ["Code", "When"],
        [
            ["200", "The engine evaluated the request (any of the five outcomes)."],
            ["404", "Unknown change_request_id or contract id."],
            ["422", "Invalid request body, or a malformed / integrity-broken upload."],
        ],
    )
    add_heading(doc, "2.3 Determinism & correlation id", 2)
    add_bullets(
        doc,
        [
            "Identical request bodies always produce a byte-identical Decision object.",
            "Responses never include the ground-truth expected_* label fields \u2014 those stay test-only.",
            "An optional X-Correlation-Id request header is echoed back in response metadata, outside the Decision object so determinism is preserved.",
        ],
    )

    add_heading(doc, "3. Endpoint summary", 1)
    add_table(
        doc,
        ["Method & Path", "Purpose"],
        [
            ["GET /health", "Liveness + loaded contract/change/rule counts."],
            ["POST /v1/decisions/by-id", "Decide a stored change request by id."],
            ["POST /v1/decisions/inline", "Decide a caller-supplied change + contract."],
            ["POST /v1/decisions/batch", "Decide many stored ids in order."],
            ["POST /v1/decisions/upload", "Score an uploaded change-request workbook."],
            ["GET /v1/contracts/{id}", "Read back a stored contract."],
            ["GET /v1/changes/{id}", "Read back a change request (labels stripped)."],
            ["GET /v1/rules", "Return the editable rulebook from rules.yaml."],
        ],
    )

    add_heading(doc, "4. Endpoint reference", 1)

    # health
    add_heading(doc, "4.1 GET /health", 2)
    add_body(doc, "Readiness probe reporting how much reference data was loaded at startup.")
    add_body(doc, "Response 200:", space_after=2)
    add_code(
        doc,
        "{\n"
        '  "status": "ok",\n'
        '  "contracts": 80,\n'
        '  "changes": 72,\n'
        '  "rules": 10\n'
        "}",
    )

    # by-id
    add_heading(doc, "4.2 POST /v1/decisions/by-id", 2)
    add_body(
        doc,
        "Resolve a change request (and its linked contract) from the loaded store "
        "and decide. Returns 404 if the id is unknown or its linked contract is "
        "missing.",
    )
    add_body(doc, "Request body:", space_after=2)
    add_code(doc, '{\n  "change_request_id": "CR-0007"\n}')
    add_body(doc, "Response 200 (DecisionResponse):", space_after=2)
    add_code(
        doc,
        "{\n"
        '  "decision": {\n'
        '    "change_request_id": "CR-0007",\n'
        '    "outcome": "APPROVE",\n'
        '    "contract_action": "Apply change to contract",\n'
        '    "winning_rule": null,\n'
        '    "explanation": "All applicable rules passed ...",\n'
        '    "tags": { "same_state": true, "signature_present": true, ... },\n'
        '    "evaluated_rules": [\n'
        '      { "rule_id": "R-001", "passed": true, "verdict": null, "message": "..." }\n'
        "    ]\n"
        "  },\n"
        '  "correlation_id": null\n'
        "}",
    )
    add_body(doc, "Optional header: X-Correlation-Id (echoed back in correlation_id).")

    # inline
    add_heading(doc, "4.3 POST /v1/decisions/inline", 2)
    add_body(
        doc,
        "Decide a caller-supplied pair instead of looking one up. The body nests "
        "the engine's ChangeRequest and ProviderContract models. Returns 422 if the "
        "body is invalid (e.g. an out-of-vocabulary enum value). The decision shape "
        "is identical to by-id.",
    )
    add_body(doc, "Request body (abridged \u2014 see \u00a75 for fields):", space_after=2)
    add_code(
        doc,
        "{\n"
        '  "change": {\n'
        '    "change_request_id": "CR-X1",\n'
        '    "linked_provider_contract_id": "PC-0001",\n'
        '    "change_category": "Practice Location Address",\n'
        '    "change_action": "Add",\n'
        '    "cms_form": "CMS-855I",\n'
        '    "provider_signature_present": "Yes",\n'
        '    "new_location_same_state": "No",\n'
        '    "development_requested": "No"\n'
        "  },\n"
        '  "contract": {\n'
        '    "provider_contract_id": "PC-0001",\n'
        '    "contract_status": "Active"\n'
        "  }\n"
        "}",
    )

    # batch
    add_heading(doc, "4.4 POST /v1/decisions/batch", 2)
    add_body(
        doc,
        "Decide many stored change requests in one call, preserving input order. "
        "Unknown ids are reported per-item (found = false) rather than failing the "
        "whole batch.",
    )
    add_body(doc, "Request body:", space_after=2)
    add_code(
        doc,
        "{\n"
        '  "change_request_ids": ["CR-0001", "CR-0002", "CR-9999"]\n'
        "}",
    )
    add_body(doc, "Response 200 (BatchResponse):", space_after=2)
    add_code(
        doc,
        "{\n"
        '  "correlation_id": null,\n'
        '  "results": [\n'
        '    { "change_request_id": "CR-0001", "found": true, "decision": { ... }, "error": null },\n'
        '    { "change_request_id": "CR-9999", "found": false, "decision": null,\n'
        '      "error": "Unknown change_request_id: CR-9999" }\n'
        "  ]\n"
        "}",
    )

    # upload
    add_heading(doc, "4.5 POST /v1/decisions/upload", 2)
    add_body(
        doc,
        "HTTP parity with the Streamlit upload+process flow. Send an .xlsx file "
        "(multipart/form-data, field name file) containing a "
        "Requested_Contract_Changes sheet. The service parses it, integrity-checks "
        "it against the reference contracts, scores every row, and returns the "
        "distribution and accuracy.",
    )
    add_body(doc, "Errors: 422 for an empty file, a malformed workbook, a missing sheet, or broken referential integrity.")
    add_body(doc, "Example with curl:", space_after=2)
    add_code(
        doc,
        "curl -X POST http://127.0.0.1:8000/v1/decisions/upload \\\n"
        '  -F "file=@data/contract_change_requests.xlsx"',
    )
    add_body(doc, "Response 200 (UploadResponse) \u2014 abridged:", space_after=2)
    add_code(
        doc,
        "{\n"
        '  "filename": "contract_change_requests.xlsx",\n'
        '  "total_changes": 72,\n'
        '  "labeled": 72,\n'
        '  "correct": 72,\n'
        '  "accuracy": 1.0,\n'
        '  "distribution": {\n'
        '    "expected": { "APPROVE": 49, "DEVELOP": 12, "DENY": 10, "INITIAL_ENROLLMENT_REQUIRED": 1 },\n'
        '    "actual":   { "APPROVE": 49, "DEVELOP": 12, "DENY": 10, "INITIAL_ENROLLMENT_REQUIRED": 1 }\n'
        "  },\n"
        '  "per_category": { "Practice Location Address": 1.0, ... },\n'
        '  "results": [ { "change_request_id": "CR-0001", "decision": { ... } } ]\n'
        "}",
    )

    # contracts / changes / rules
    add_heading(doc, "4.6 GET /v1/contracts/{id}", 2)
    add_body(doc, "Read back a stored provider contract by id. Returns 404 if missing.")
    add_heading(doc, "4.7 GET /v1/changes/{id}", 2)
    add_body(
        doc,
        "Read back a change request by id, with the ground-truth label fields "
        "(expected_validation_outcome, expected_contract_action) stripped. Returns "
        "404 if missing.",
    )
    add_heading(doc, "4.8 GET /v1/rules", 2)
    add_body(doc, "Return the editable rulebook exactly as stored in config/rules.yaml.")

    add_heading(doc, "5. Data models", 1)
    add_heading(doc, "5.1 ChangeRequest (selected fields)", 2)
    add_table(
        doc,
        ["Field", "Type", "Notes"],
        [
            ["change_request_id", "str", "Primary key."],
            ["linked_provider_contract_id", "str", "FK to a contract (must resolve)."],
            ["change_category", "enum", "e.g. Practice Location Address, EFT Information."],
            ["change_action", "enum", "Add | Change | Terminate."],
            ["cms_form", "str", "Must be CMS-855I to pass R-001."],
            ["provider_signature_present", "enum", "Yes | No | N/A (R-008)."],
            ["new_location_same_state", "enum", "Yes | No | N/A (drives R-003)."],
            ["development_requested", "enum", "Yes | No | N/A (correlates with DEVELOP)."],
            ["requested_new_value", "str?", "The proposed value for the change."],
            ["expected_validation_outcome", "enum?", "Label \u2014 never returned by the API."],
        ],
    )
    add_heading(doc, "5.2 ProviderContract (selected fields)", 2)
    add_table(
        doc,
        ["Field", "Type", "Notes"],
        [
            ["provider_contract_id", "str", "Primary key."],
            ["npi / ptan", "str?", "Identifiers matched by R-002."],
            ["legal_business_name", "str?", "Matched by R-002."],
            ["contract_status", "enum?", "Active | Suspended."],
            ["final_adverse_action_on_file", "str?", "Adverse-action context."],
            ["sanction_screening_status", "str?", "Screening state (baseline 'Clear')."],
        ],
    )
    add_body(
        doc,
        "The full models map all 43 contract columns and 25 change-request columns; "
        "the long tail is optional. See engine/models.py for the complete definition.",
    )
    add_heading(doc, "5.3 Decision (response payload)", 2)
    add_table(
        doc,
        ["Field", "Type", "Notes"],
        [
            ["outcome", "enum", "One of the five outcomes."],
            ["contract_action", "str", "Canonical downstream action for the outcome."],
            ["winning_rule", "str?", "The rule that decided a non-APPROVE outcome."],
            ["explanation", "str", "Human-readable rationale."],
            ["tags", "object", "Derived boolean/enum signals that fired."],
            ["evaluated_rules", "list", "Each applicable rule + passed/verdict/message."],
        ],
    )

    add_heading(doc, "6. The rulebook (R-001 \u2013 R-010)", 1)
    add_body(
        doc,
        "Only rules whose applies_to_change_category matches the request (or 'All') "
        "fire. Each returns a pass/fail with a failure outcome:",
    )
    add_table(
        doc,
        ["Rule", "Checks", "On failure"],
        [
            ["R-001", "CMS form is CMS-855I", "REJECT"],
            ["R-002", "NPI/PTAN/TIN/LBN match the contract", "DEVELOP"],
            ["R-003", "Not an out-of-state practice-location add", "INITIAL_ENROLLMENT_REQUIRED"],
            ["R-004", "Special payment address is an allowed type", "DEVELOP"],
            ["R-005", "No exclusion/debarment/revocation", "DENY"],
            ["R-006", "Reassignment target is enrolled", "DEVELOP"],
            ["R-007", "EFT change has CMS-588 + bank docs", "DEVELOP"],
            ["R-008", "Certification statement is signed", "REJECT"],
            ["R-009", "Specialty compatible with provider type", "DENY"],
            ["R-010", "Existing docs reused unless new evidence needed", "DEVELOP"],
        ],
    )

    add_heading(doc, "7. Decision precedence ladder", 1)
    add_body(
        doc,
        "decide() evaluates the rules, then resolves a single outcome via an "
        "ordered ladder \u2014 first match wins, then it stops:",
    )
    add_table(
        doc,
        ["Tier", "Condition", "Outcome"],
        [
            ["1", "Exclusion/debarment (R-005) or incompatible specialty (R-009)", "DENY"],
            ["2", "Wrong CMS form (R-001) or missing signature (R-008)", "REJECT"],
            ["3", "Out-of-state practice-location add (R-003)", "INITIAL_ENROLLMENT_REQUIRED"],
            ["4", "Any DEVELOP-class rule failed / docs incomplete", "DEVELOP"],
            ["5", "Default", "APPROVE"],
        ],
    )
    add_body(
        doc,
        "Tier 1 outranks tier 2: an adverse request that is also missing a signature "
        "still returns DENY.",
    )

    add_heading(doc, "8. Configuration", 1)
    add_table(
        doc,
        ["Variable", "Effect"],
        [
            ["DATA_PATH", "Directory with reference_data.xlsx + contract_change_requests.xlsx. Unset = bundled data/."],
            ["ALLOWED_ORIGINS", "Comma-separated CORS origins. Unset = none."],
        ],
    )

    add_heading(doc, "9. Quick start", 1)
    add_code(
        doc,
        "uvicorn api.main:app --reload\n"
        "# then open http://127.0.0.1:8000/docs",
    )
    add_body(
        doc,
        "See the companion Step-by-Step Local Setup Guide for full environment "
        "setup. The interactive Swagger UI at /docs lets you exercise every "
        "endpoint without writing any client code.",
    )

    out = DOCS / "API_Handbook.docx"
    doc.save(out)
    return out


def main():
    DOCS.mkdir(exist_ok=True)
    p1 = build_local_setup()
    p2 = build_api_handbook()
    print(f"Wrote: {p1}")
    print(f"Wrote: {p2}")


if __name__ == "__main__":
    main()
